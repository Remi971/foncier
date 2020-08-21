# coding: utf-8

import csv
import pandas as pd
import geopandas as gpd
from docx import Document
from time import strftime, localtime
#from docx.shared import Inches POUR les images
donnees = {
    "gpkg": {
        "nomGPKG": {},
        "layers": {}
        },
    "dossier": {
        "chemin": "C:/Users/Citadia/Dev/projet_foncier/foncier/data",
        "couches": {"Bâti": "BATI_region.shp", "Parcelles": "PARCELLE_region.shp", "Structuration territoriale": "ZONE U PLUi_ENNEZAT.shp"}
        },
    "paramètres": {
        "défauts": "vide",
        "perso": {
            "champs": "LIBELLE",
            "valeurs": {
                "UAa": {"non-batie": 400, "batie": 1000, "cesMax": 40, "test": 10, "bufBati": 8},
                "UAi": {"non-batie": 400, "batie": 1000, "cesMax": 40, "test": 10, "bufBati": 8},
                "UCv": {"non-batie": 400, "batie": 1000, "cesMax": 40, "test": 10, "bufBati": 8},
                "UE": {"non-batie": 400, "batie": 1000, "cesMax": 40, "test": 10, "bufBati": 8},
                "UG": {"non-batie": 400, "batie": 1000, "cesMax": 40, "test": 10, "bufBati": 8},
                "UJ": {"non-batie": 400, "batie": 1000, "cesMax": 40, "test": 10, "bufBati": 8},
                "UR": {"non-batie": 400, "batie": 1000, "cesMax": 40, "test": 10, "bufBati": 8}
                }
            },
        "filtres": {}}}
dossier = 'C:/Users/Citadia/Dev/projet_foncier/foncier'
date = '210820'
def exportReglages(data, dossier, date2):
    document = Document()
    # TITRE
    document.add_heading('Réglages', 0)
    date = document.add_paragraph(f'Date : {strftime("%a, %d %b %Y %H:%M:%S", localtime())}')
    #Nom de la source de donnée et noms des couches utilisées
    if data["gpkg"]["nomGPKG"] != {}:
        document.add_heading(f'Base de données Geopackage : {data["gpkg"]["nomGPKG"]}', level=2)
        for layer in data["gpkg"]["layers"]:
            document.add_paragraph(f'{layer} : {data["gpkg"]["layers"][layer]}', style='List Bullet')

    if data["dossier"]["chemin"] != {}:
        document.add_heading(f'Base de données dans le dossier : {data["dossier"]["chemin"]}', level=2)
        for couche in data["dossier"]["couches"]:
            document.add_paragraph(f'{couche} : {data["dossier"]["couches"][couche]}', style='List Bullet')
    #Paramètres
    document.add_heading('Paramètres', 1)
    table = document.add_table(rows=1, cols=6)
    table.style = 'LightGrid-Accent1'
    hdr_cells = table.rows[0].cells
    if data["paramètres"]["perso"] != 'vide':
        hdr_cells[0].text = data["paramètres"]["perso"]["champs"]
    else:
        hdr_cells[0].text ='Enveloppe'
    hdr_cells[1].text = "Surface minimal de la parcelle non bâtie"
    hdr_cells[2].text = "Surface minimale de la parcelle bâtie"
    hdr_cells[3].text = "CES maximum de la parcelle divisible"
    hdr_cells[4].text = "Distance du buffer pour le test"
    hdr_cells[5].text = "Distance du buffer autour du bâti"
    #import pdb; pdb.set_trace()
    if data["paramètres"]["perso"] != 'vide':
        for k, v in data["paramètres"]["perso"]["valeurs"].items():
            row_cells = table.add_row().cells
            row_cells[0].text = k
            row_cells[1].text = str(v["non-batie"])
            row_cells[2].text = str(v["batie"])
            row_cells[3].text = str(v["cesMax"])
            row_cells[4].text = str(v["test"])
            row_cells[5].text = str(v["bufBati"])
    else:
        row_cells = table.add_row().cells
        row_cells[0].text = 'Enveloppe'
        row_cells[1].text = str(data["paramètres"]["défauts"]["non-batie"])
        row_cells[2].text = str(data["paramètres"]["défauts"]["batie"])
        row_cells[3].text = str(data["paramètres"]["défauts"]["cesMax"])
        row_cells[4].text = str(data["paramètres"]["défauts"]["test"])
        row_cells[5].text = str(data["paramètres"]["défauts"]["bufBati"])

    nom = f'{dossier}/{date2}_reglages.docx'
    document.save(nom)

def export_reglages_csv(data, dossier, date):
    nom = f'{dossier}/{date}_reglages.csv'
    with open(nom, 'w', newline='') as csvfile:
        fieldname = ['Source','non-batie', 'batie', 'cesMax', 'test', 'bufBati']
        thewriter = csv.DictWriter(csvfile, fieldnames=fieldname, delimiter=';',
                                quotechar = '|', quoting=csv.QUOTE_MINIMAL)
        thewriter.writeheader()
        if data["paramètres"]["défauts"] == "vide":
            for k, v in data["paramètres"]["perso"]["valeurs"].items():
                thewriter.writerow({'Source' : k,
                                    'non-batie' : v['non-batie'],
                                    'batie' : v['batie'],
                                    'cesMax' : v['cesMax'],
                                    'test' : v['test'],
                                    'bufBati' : v['bufBati'],
                                     })
        else:
            thewriter.writerow({
            'Source' : "Par défauts",
            'non-batie': data["paramètres"]["défauts"]["non-batie"],
            'batie': data["paramètres"]["défauts"]["batie"],
            'cesMax': data["paramètres"]["défauts"]["cesMax"],
            'test': data["paramètres"]["défauts"]["test"],
            'bufBati': data["paramètres"]["défauts"]["bufBati"],
            })

exportReglages(donnees, dossier,date)
export_reglages_csv(donnees, dossier,date)
# def export_table(data):
#     table = {}
#     if data["paramètres"]["défauts"] == "vide":
#         for k, v in data["paramètres"]["perso"]["valeurs"].items():
#             table['Source'] = k,
#             table['non-batie'] = v['non-batie']
#             table['batie'] : v['batie']
#             table['cesMax'] : v['cesMax']
#             table['test'] : v['test']
#             table['bufBati'] : v['bufBati']
#     else:
#         table['Source'] : "Par défauts"
#         table['non-batie']: data["paramètres"]["défauts"]["non-batie"]
#         table['batie']: data["paramètres"]["défauts"]["batie"]
#         table['cesMax']: data["paramètres"]["défauts"]["cesMax"]
#         table['test']: data["paramètres"]["défauts"]["test"]
#         table['bufBati']: data["paramètres"]["défauts"]["bufBati"]
#     return table
